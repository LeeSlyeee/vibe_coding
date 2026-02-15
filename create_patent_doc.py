from docx import Document
from docx.shared import Pt
import os

def create_patent_doc():
    doc = Document()
    
    # Title
    doc.add_heading('【고안의 설명】', level=0)
    
    # 1. 고안의 명칭
    doc.add_heading('【고안의 명칭】', level=1)
    p = doc.add_paragraph()
    run = p.add_run('온디바이스 인공지능을 이용한 프라이버시 보장형 멘탈케어 서비스 제공 시스템 및 그 방법\n(System and method for providing privacy-guaranteed mental care service using on-device artificial intelligence)')
    
    # 2. 기술분야
    doc.add_heading('【기술분야】', level=1)
    doc.add_paragraph('본 고안은 디지털 헬스케어 및 인공지능 기술 분야에 관한 것으로, 구체적으로는 사용자의 단말기 내에서 독립적으로 구동되는 경량화 언어 모델(sLLM)을 이용하여 민감한 심리/의료 데이터의 외부 유출 없이 멘탈케어 서비스를 제공하고, 의료진에게는 필요한 임상 데이터만을 암호화하여 동기화하는 온디바이스 AI 기반 멘탈케어 시스템 및 방법에 관한 것이다.')

    # 3. 고안의 배경이 되는 기술
    doc.add_heading('【고안의 배경이 되는 기술】', level=1)
    doc.add_paragraph('현대 사회에서 우울증, 불안 장애 등 정신 건강 문제의 관리 필요성이 대두됨에 따라 다양한 모바일 멘탈케어 및 웰니스 애플리케이션이 개발되고 있다. 기존의 서비스들은 주로 사용자의 입력 데이터를 중앙 클라우드 서버로 전송하여 거대 언어 모델(LLM)을 통해 분석하는 방식을 취한다.')
    doc.add_paragraph('그러나 이러한 서버 기반(Server-based) 또는 클라우드 기반 처리 방식은 사용자의 지극히 사적인 심리 상태나 의료 기록이 외부 네트워크를 통해 전송되고 제3의 서버에 저장된다는 점에서 심각한 프라이버시 침해 우려를 낳는다. 특히, 정신과 진료 기록이나 심리 상담 내용은 유출 시 사회적 낙인(Stigma) 효과를 유발할 수 있어 보안이 무엇보다 중요함에도 불구하고, 기존 기술은 중앙 집중식 처리의 한계로 인해 데이터 주권(Data Sovereignty)을 사용자에게 온전히 보장하지 못하는 문제점이 있다.')

    # 4. 선행기술문헌
    doc.add_heading('【선행기술문헌】', level=1)
    
    doc.add_heading('(【특허문헌】)', level=2)
    doc.add_paragraph('1. 공개특허 (출원인: 주식회사 마인드트리)\n   - 명칭: 인공지능 기반 상담사-내담자 매칭 시스템 및 심리데이터 통합 분석을 통한 맞춤형 상담 지원 방법\n   - 내용: EAP(근로자 지원 프로그램) 및 AI를 활용하여 상담사와 내담자를 매칭하고, 상담 데이터를 분석하여 맞춤형 상담을 지원하는 기술이다. 그러나 이는 상담사와 서버 간의 데이터 전송을 전제로 하며, 디바이스 자체에서의 완결성(On-Device)보다는 매칭 및 분석 효율화에 초점이 맞춰져 있다.')
    doc.add_paragraph('2. 공개특허 (출원인: 주식회사 로완)\n   - 명칭: LLM(거대언어모델)을 이용한 디지털 치료제 콘텐츠 추천 방법 및 장치 (2023년 6월 출원)\n   - 내용: 초거대 언어 모델을 활용하여 사용자의 인지 기능 훈련 및 치료 콘텐츠를 추천하는 기술이다. 서버의 강력한 연산 능력을 활용하는 장점이 있으나, 민감한 개인 의료 데이터가 외부 서버로 전송되어 처리되는 구조적 한계를 가진다.')
    doc.add_paragraph('3. 공개특허 (출원인: 디퍼아이)\n   - 명칭: 엣지 AI 반도체 및 데이터 경량화 처리 기술 관련 특허\n   - 내용: 엣지 컴퓨팅 환경에서의 데이터 처리 및 저전력 구동에 관한 하드웨어 및 데이터 처리 기술로, 본 고안의 기술적 구현(On-Device)과 관련된 선행 기술이다.')

    doc.add_heading('(【비특허문헌】)', level=2)
    doc.add_paragraph('1. Mustaqeem Khan et al., "Edge-AI Surveillance System for Emotion Recognition in Smart Cities", 202X. (엣지 컴퓨팅 기반의 감정 인식 효용성을 다룬 학술 논문)')
    doc.add_paragraph('2. 유럽연합(EU) AI 법(AI Act) 및 GDPR 가이드라인 (정신건강 등 고위험 AI에 대한 데이터 보안 및 투명성 규제 관련 문헌)')

    # 5. 고안의 내용
    doc.add_heading('【고안의 내용】', level=1)
    
    doc.add_heading('【해결하고자 하는 과제】', level=2)
    doc.add_paragraph('본 고안은 상기 선행기술들이 가지는 중앙 서버 의존성 및 데이터 보안 취약점을 해결하기 위해 안출된 것으로,')
    doc.add_paragraph('첫째, 사용자의 민감한 상담 내용과 생체 데이터를 외부 서버로 전송하지 않고 사용자 단말기(Edge Device) 내부에서 AI 추론을 수행함으로써 데이터 유출 위험을 원천 차단하는 것을 목적으로 한다.')
    doc.add_paragraph('둘째, 인터넷 연결이 없는 환경(Offline)에서도 끊김 없는 멘탈케어 서비스를 제공하며, 대규모 서버 구축 비용을 절감하는 것을 목적으로 한다.')
    doc.add_paragraph('셋째, 사용자 단말에서 분석된 결과 중 임상적으로 유의미한 데이터만을 선별 및 암호화하여 의료진 전용 대시보드와 안전하게 동기화함으로써, 진료 효율성을 높이고 예후 관리를 돕는 것을 목적으로 한다.')

    doc.add_heading('【과제의 해결 수단】', level=2)
    doc.add_paragraph('상기 과제를 해결하기 위한 본 고안의 시스템은,')
    doc.add_paragraph('사용자의 음성 또는 텍스트 데이터를 입력받는 입력부;')
    doc.add_paragraph('상기 입력된 데이터를 분석하기 위해 사용자 단말기에 탑재된 온디바이스(On-Device) 경량화 언어 모델(sLLM, 예: Gemma-2-2b 등)을 포함하는 로컬 AI 추론 엔진;')
    doc.add_paragraph('상기 추론 엔진의 분석 결과를 바탕으로 상담 피드백 및 감정 온도(Emotional Temperature)를 생성하여 사용자에게 제공하는 출력부;')
    doc.add_paragraph('생성된 상담 데이터 및 분석 데이터를 단말기 내의 보안 영역에 저장하는 로컬 데이터베이스; 및')
    doc.add_paragraph('기 설정된 의료진 단말 또는 병원 서버와 인증된 보안 채널을 형성하고, 상기 로컬 데이터베이스에 저장된 데이터 중 사용자 동의에 의해 선택된 임상 지표(위험 징후, 감정 추이 등)만을 암호화하여 동기화하는 데이터 동기화 모듈;을 포함하는 것을 특징으로 한다.')

    doc.add_heading('【고안의 효과】', level=2)
    doc.add_paragraph('본 고안에 따르면 다음과 같은 효과가 있다.')
    doc.add_paragraph('1. 데이터 보안 및 프라이버시 강화: 민감한 심리 데이터가 사용자 단말을 벗어나지 않고 처리되므로, 개인정보보호법 및 의료법(HIPAA 등) 준수에 탁월하며 사용자의 심리적 장벽을 낮춘다.')
    doc.add_paragraph('2. 비용 절감 및 효율성: 고가의 중앙 GPU 서버 비용을 절감하고, 사용자 단말의 컴퓨팅 파워를 활용하는 엣지 컴퓨팅 기술로 서비스 유지 비용을 획기적으로 낮춘다.')
    doc.add_paragraph('3. 진료 연속성 확보: 의료진은 환자의 내원 간격 사이에 발생한 구체적인 감정 변화와 위험 신호(Red Flag)를 대시보드를 통해 정량적으로 파악할 수 있어 진료의 질을 높일 수 있다.')
    doc.add_paragraph('4. 사용 편의성: 네트워크 연결 여부와 관계없이 언제 어디서나 즉각적인 AI 상담 서비스 이용이 가능하다.')

    # 6. 도면의 간단한 설명
    doc.add_heading('【도면의 간단한 설명】', level=1)
    doc.add_paragraph('제1도는 본 고안의 일 실시예에 따른 온디바이스 AI 멘탈케어 시스템의 전체적인 블록 구성도이다.')
    doc.add_paragraph('제2도는 사용자 단말(App)과 의료진 대시보드(Web) 간의 데이터 흐름 및 보안 동기화 과정을 나타낸 순서도이다.')
    doc.add_paragraph('제3도는 온디바이스 AI 엔진이 사용자 입력을 처리하여 감정 온도를 산출하고 피드백을 생성하는 로직 흐름도이다.')
    doc.add_paragraph('제4도는 의료진에게 제공되는 환자 모니터링 대시보드의 예시 화면이다.')

    # 7. 고안을 실시하기 위한 구체적인 내용
    doc.add_heading('【고안을 실시하기 위한 구체적인 내용】', level=1)
    doc.add_paragraph('이하, 첨부된 도면을 참조하여 본 고안의 바람직한 실시예를 상세히 설명한다.')
    doc.add_paragraph('본 고안의 시스템은 크게 환자용 모바일 애플리케이션(Client App)과 의료진용 웹 대시보드(Provider Web)로 구성된다.')
    
    doc.add_paragraph('1. 환자용 모바일 애플리케이션 (HaruON App)')
    doc.add_paragraph('   - 입력 인터페이스: 사용자는 텍스트 또는 음성(STT)을 통해 그날의 감정이나 겪었던 일(Event)을 기록한다.')
    doc.add_paragraph('   - 로컬 LLM 엔진: 앱 내부에는 경량화된 언어 모델(예: Gemma 2 2b, Llama 3 8b Quantized 등)이 탑재되어 있다. 이 엔진은 네트워크 통신 없이 사용자의 입력을 실시간으로 분석하여, 감정의 종류와 강도를 수치화(감정 온도)하고 위로와 공감의 피드백 메시지를 생성한다.')
    doc.add_paragraph('   - 로컬 저장소: 분석된 원본 데이터(Raw Data)는 사용자의 스마트폰 내 샌드박스(Sandbox) 처리된 로컬 DB에만 저장되며, 외부로 전송되지 않는다.')

    doc.add_paragraph('2. 데이터 동기화 및 보안 전송')
    doc.add_paragraph('   - 앱은 의료진 서버와 동기화 시, 전체 대화 로그가 아닌 의료적으로 유의미한 요약 데이터(Summary), 감정 변화 추이, 위험 키워드(자살 사고, 자해 위협 등)만을 추출한다.')
    doc.add_paragraph('   - 추출된 데이터는 종단간 암호화(E2EE) 기술을 통해 병원 내 로컬 서버(On-Premise) 또는 보안 클라우드로 전송된다.')

    doc.add_paragraph('3. 의료진용 웹 대시보드')
    doc.add_paragraph('   - 의료진은 대시보드를 통해 환자들의 일별 감정 온도 변화 그래프, 수면 시간, 스트레스 지수 등을 한눈에 모니터링한다.')
    doc.add_paragraph('   - 특정 환자에게서 "죽음", "포기" 등의 위험 단어가 감지될 경우, 시스템은 의료진에게 즉시 "Red Flag" 알림을 발송하여 조기 개입을 유도한다.')

    # 8. 실시예
    doc.add_heading('(【실시예】)', level=1)
    doc.add_paragraph('[가상의 실시예 1] 우울증 환자에 대한 적용 시나리오')
    doc.add_paragraph('(본 실시예는 시스템의 동작 이해를 돕기 위한 가상의 사용자 시나리오임)')
    doc.add_paragraph('우울증 진단을 받은 사용자 A가 스마트폰이 "비행기 모드"인 상태에서 "오늘 너무 힘들어서 아무것도 하기 싫어"라고 앱에 기록한다.')
    doc.add_paragraph('-> 앱 내 탑재된 AI 모델이 즉시 이를 분석하여 "무기력감이 크게 느껴지시는군요. 잠시 따뜻한 물 한 잔 마시는 건 어떨까요?"라는 피드백을 출력한다.')
    doc.add_paragraph('-> 동시에 내부적으로 "무기력 지수: 85", "감정 온도: 10도(매우 낮음)"으로 기록한다.')
    doc.add_paragraph('-> 추후 네트워크가 연결되면, 해당 수치 데이터만 담당 주치의의 대시보드로 전송되어 주치의가 다음 진료 시 A의 상태 악화를 인지하고 상담에 활용한다.')

    # 9. 산업상 이용가능성
    doc.add_heading('(【산업상 이용가능성】)', level=1)
    doc.add_paragraph('본 고안은 정신건강의학과 병·의원, 보건소, 정신건강복지센터의 환자 관리 시스템으로 이용 가능하며, 일반 기업의 임직원 멘탈케어(EAP) 서비스 등 B2B, B2G 헬스케어 시장 전반에서 산업적으로 이용 가능하다.')

    # 10. 부호의 설명
    doc.add_heading('(【부호의 설명】)', level=1)
    doc.add_paragraph('100: 사용자 단말기 (스마트폰)')
    doc.add_paragraph('110: 입력부 (마이크/키보드)')
    doc.add_paragraph('120: 온디바이스 AI 엔진 (sLLM)')
    doc.add_paragraph('130: 로컬 암호화 저장소')
    doc.add_paragraph('200: 네트워크 (보안 채널)')
    doc.add_paragraph('300: 의료진 관리 서버')
    doc.add_paragraph('310: 데이터 동기화 모듈')
    doc.add_paragraph('320: 의료진 대시보드 (Admin Web)')

    # 11. 수탁번호
    doc.add_heading('(【수탁번호】)', level=1)
    doc.add_paragraph('해당사항 없음')

    # Save the file
    file_path = 'patent_application_draft.docx'
    doc.save(file_path)
    print(f"File created successfully at: {os.path.abspath(file_path)}")

if __name__ == "__main__":
    create_patent_doc()
