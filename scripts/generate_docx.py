#!/usr/bin/env python3
"""
2025년도 예비창업패키지 사업계획서 DOCX 생성 스크립트 v2
- 원본 양식 복사
- 파란색 예시 답변 삭제
- 마음온(Maum-On) 프로젝트 내용으로 채우기

TABLE 구조 (확인 완료):
  T0: 안내 문구 (삭제할 목차 페이지 안내)
  T1: 목차 테이블 (항목 / 세부항목)
  T2: 세부 목차 (1,2,3,4)
  T3: 개요(요약) 테이블 — 8 rows, 셀 병합 복잡
    R0: [명칭(gs1)][값(gs1)][범주(gs2)][값(gs1)]    -> 4 xml tcs
    R1~R5: [제목(gs1)][내용(gs4)]                     -> 2 xml tcs
    R6: [이미지(gs1)][img_area(gs2)][img_area(gs2)]   -> 3 xml tcs
    R7: [이미지(gs1)][제목(gs2)][제목(gs2)]            -> 3 xml tcs
  T4: 사업추진 일정(협약기간 내) — 6 rows
  T5: 1단계 예산 — 7 rows (R0 헤더, R1~R5 항목, R6 합계)
  T6: 2단계 예산 — 7 rows (동일 구조)
  T7: 사업추진 일정(전체) — 6 rows
  T8: 팀 구성(안) — 4 rows
  T9: 협력 기관 — 4 rows
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

SRC = "[별첨 1] 2025년도 예비창업패키지 사업계획서 양식.docx"
DST = "2025_예비창업패키지_사업계획서_마음온.docx"

BLACK = RGBColor(0, 0, 0)
doc = Document(SRC)


# ───────────────── HELPER ─────────────────

def set_cell_text(cell, text, bold=False, font_size=10, color=BLACK):
    """셀 전체를 비우고 단일 텍스트로 설정"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    p = cell.paragraphs[0]
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color


def set_tc_text(tc, text, bold=False, font_size=10, color=BLACK):
    """XML tc 엘리먼트의 텍스트를 직접 설정 (병합된 셀 처리용)"""
    # 기존 run 텍스트 비우기
    for r in tc.findall(f'.//{qn("w:r")}'):
        for t in r.findall(qn("w:t")):
            t.text = ""
    # 첫 번째 run을 찾아 텍스트 설색, 없으면 p/r/t 생성
    runs = tc.findall(f'.//{qn("w:r")}')
    if runs:
        ts = runs[0].findall(qn("w:t"))
        if ts:
            ts[0].text = text
        else:
            from lxml import etree
            t_el = etree.SubElement(runs[0], qn("w:t"))
            t_el.text = text
    else:
        # 첫 p에 새 r/t 추가
        ps = tc.findall(qn("w:p"))
        if not ps:
            from lxml import etree
            p_el = etree.SubElement(tc, qn("w:p"))
            ps = [p_el]
        from lxml import etree
        r_el = etree.SubElement(ps[0], qn("w:r"))
        rPr = etree.SubElement(r_el, qn("w:rPr"))
        sz = etree.SubElement(rPr, qn("w:sz"))
        sz.set(qn("w:val"), str(font_size * 2))  # half-points
        if bold:
            etree.SubElement(rPr, qn("w:b"))
        color_el = etree.SubElement(rPr, qn("w:color"))
        color_el.set(qn("w:val"), f"{color}")
        t_el = etree.SubElement(r_el, qn("w:t"))
        t_el.text = text
        t_el.set(qn("xml:space"), "preserve")


def set_para_heading(para, title, font_size=14):
    """ㅇ + 볼드 제목 형식으로 설정"""
    for run in para.runs:
        run.text = ""
    r1 = para.runs[0] if para.runs else para.add_run()
    r1.text = "ㅇ "
    r1.bold = False
    r1.font.size = Pt(font_size)
    r1.font.color.rgb = BLACK
    r2 = para.add_run(title)
    r2.bold = True
    r2.font.size = Pt(font_size)
    r2.font.color.rgb = BLACK


def set_para_body(para, text, font_size=10):
    """본문 텍스트 설정"""
    for run in para.runs:
        run.text = ""
    r = para.runs[0] if para.runs else para.add_run()
    r.text = text
    r.bold = False
    r.font.size = Pt(font_size)
    r.font.color.rgb = BLACK


def clear_blue_globally():
    """문서 전체에서 파란색(0000FF, 001AFF) run의 텍스트를 제거"""
    blue_colors = {'0000FF', '001AFF'}
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color and run.font.color.rgb and str(run.font.color.rgb) in blue_colors:
                run.text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.color and run.font.color.rgb and str(run.font.color.rgb) in blue_colors:
                            run.text = ""


# ═══════════════════════════════════════════════════════
# STEP 0: 파란색 가이드 텍스트 전역 삭제
# ═══════════════════════════════════════════════════════
clear_blue_globally()
print("✅ 파란색 가이드 텍스트 삭제 완료")


# ═══════════════════════════════════════════════════════
# TABLE 3: 개요(요약) — 셀 병합을 고려한 XML 직접 접근
# ═══════════════════════════════════════════════════════
t3 = doc.tables[3]

# R0: 명칭 / 범주 (4 xml tcs)
r0_tcs = t3.rows[0]._tr.findall(qn('w:tc'))
set_tc_text(r0_tcs[1], "마음온(Maum-On)", bold=True, font_size=10)
set_tc_text(r0_tcs[3], "디지털 헬스케어 / AI 정신건강 관리 플랫폼", font_size=9)

# R1~R5: 각 2 xml tcs (제목 | 내용gs4)
summaries = {
    1: (  # 아이템 개요
        "마음온(Maum-On)은 On-Device AI(경량 언어모델, sLLM)를 활용하여 사용자의 스마트폰 내에서 "
        "모든 감정 데이터를 처리하는 프라이버시 퍼스트(Privacy-First) 정신건강 관리 플랫폼이다. "
        "\"Network-OFF, AI-ON\"이라는 슬로건 아래, 사용자의 민감한 일기·감정 데이터가 서버로 전송되지 "
        "않으면서도 24시간 AI 감정 상담을 제공한다. 동시에, 암호화된 감정 트렌드와 위기 알림(Red Flag Alert)을 "
        "의료기관 대시보드에 연동하여, 진료 사이의 '임상적 사각지대'를 해소하는 B2B/B2G 임상 연동 생태계를 구축한다."
    ),
    2: (  # 문제 인식
        "정신건강 관리에서 ①개인정보 유출에 대한 공포(Privacy Paradox), ②클라우드 AI의 높은 운영 비용, "
        "③진료 간 환자 상태 파악 불가(임상적 사각지대)라는 3대 문제가 존재한다. MZ세대를 중심으로 "
        "정신건강에 대한 관심이 폭발적으로 증가하고 있으나, 기존 AI 상담 서비스는 대화 내용이 외부 서버에 "
        "저장·학습되는 구조적 한계로 인해 사용자의 솔직한 감정 표현을 억제하고 있다."
    ),
    3: (  # 실현 가능성
        "Gemma-2-2b를 4bit 양자화하여 아이폰에서 구동 가능한 On-Device sLLM을 개발 완료하였으며, "
        "Apple MLX 프레임워크를 통한 최적화 파이프라인을 구축하였다. 서버사이드에서는 Llama-3-8B 기반의 "
        "심층 분석 엔진과 RunPod Serverless GPU 인프라로 확장성을 확보하였다. Django + Vue.js + SwiftUI "
        "풀스택 아키텍처와 Fernet 대칭키 암호화를 통한 보안 체계가 이미 구현되어 임상 파일럿이 가능한 수준이다."
    ),
    4: (  # 성장전략
        "Phase 1: 대학병원/정신건강복지센터 파일럿을 통한 레퍼런스 확보 → Phase 2: 디지털 치료제(DTx) 인증 "
        "및 보험 수가 적용 추진 → Phase 3: 전국 보건소 B2G 확대 및 기업용 B2B SDK 라이선싱. 수익 모델은 "
        "기관 라이선스(SaaS) 기반의 안정적 B2B/B2G 모델이며, 경쟁사 대비 '데이터 주권(On-Device)'과 "
        "'임상 연동'이라는 차별화된 포지셔닝을 확보한다."
    ),
    5: (  # 팀 구성
        "대표자는 풀스택 개발(Vue.js, Swift, Python/Django) 및 AI 엔지니어링(모델 양자화, 파인튜닝) 역량을 "
        "보유하며, 프로토타입 전 과정을 독자 완수한 실적이 있다. 임상심리학 박사 및 정신건강의학 전문의 "
        "자문위원을 확보하였으며, 협약기간 내 프론트엔드 개발자 및 UX/UI 디자이너 채용을 계획하고 있다."
    ),
}

for row_idx, text in summaries.items():
    tcs = t3.rows[row_idx]._tr.findall(qn('w:tc'))
    # 두 번째 tc (gs=4)가 내용 영역
    if len(tcs) >= 2:
        set_tc_text(tcs[1], text, font_size=9)

# R7: 이미지 제목 (3 xml tcs: [이미지(gs1)][제목1(gs2)][제목2(gs2)])
r7_tcs = t3.rows[7]._tr.findall(qn('w:tc'))
if len(r7_tcs) >= 3:
    set_tc_text(r7_tcs[1], "iOS 앱 메인 — 마음 온도 / AI 감정 상담 다이어리", font_size=8)
    set_tc_text(r7_tcs[2], "임상 모니터링 대시보드 / 시스템 아키텍처", font_size=8)

print("✅ TABLE 3 개요 요약 완료")


# ═══════════════════════════════════════════════════════
# 본문 단락 채우기
# ═══════════════════════════════════════════════════════

# ── 1. 문제 인식 (Problem) ──
para_contents = [
    # (idx, type, content)
    (17, 'heading', "국내·외 시장 현황 및 문제점"),
    (19, 'body', (
        "- 2024년 기준 국내 정신건강 관련 디지털 헬스케어 시장 규모는 약 1조 원으로 추산되며, "
        "연평균 15% 이상의 성장세를 보이고 있다. 특히 COVID-19 이후 MZ세대를 중심으로 정신건강에 "
        "대한 관심이 폭발적으로 증가하였으나, 전문 인력 부족과 접근성 문제로 인해 적절한 관리를 "
        "받지 못하는 사각지대가 광범위하게 존재한다.\n"
        "- 기존 AI 상담 서비스(ChatGPT, 마인디 등)는 사용자 대화를 클라우드 서버에 전송·저장하여 "
        "민감한 정신건강 데이터의 유출·학습 위험이 상존한다. 이로 인해 사용자가 솔직한 감정 표현을 "
        "자기 검열하여, AI 상담의 핵심 가치인 '자기 노출(Self-Disclosure)' 효과가 반감된다.\n"
        "- 정신건강 의료기관에서는 환자의 진료 사이(2~4주) 일상 속 감정 변화를 파악할 방법이 부재하여 "
        "위기 상황(자해/자살 충동)의 조기 발견이 어렵다.\n"
        "- 글로벌 디지털 정신건강 시장은 2025년 약 $17.5B 규모로 예상되며, B2B/B2G 임상 연동 및 "
        "On-Device AI 기반 프라이버시 보장 솔루션은 사실상 부재한 상황이다."
    )),
    (21, 'heading', "개발 필요성"),
    (23, 'body', (
        "- 개인정보보호법 강화 및 마이데이터 사업 확대에 따라, 정신건강 데이터의 '데이터 주권' "
        "보장은 선택이 아닌 필수가 되고 있다. On-Device AI를 통해 데이터가 기기를 떠나지 않는 "
        "구조는 법적·윤리적 요구사항을 기술적으로 해결하는 유일한 접근법이다.\n"
        "- 정신건강복지센터, 보건소 등 공공기관은 디지털 도구 도입을 추진 중이나, 보안 요건이 높아 "
        "기존 클라우드 방식의 채택이 어렵다. On-Device 처리 + 암호화 연동 방식은 기관의 보안 기준을 "
        "충족하면서도 실시간 모니터링을 가능하게 한다.\n"
        "- 기존 서비스는 사용자 1인당 월 $0.5~$2.0의 API 토큰 비용이 발생하나, On-Device AI는 "
        "API 토큰 비용 제로(Zero Token Cost)를 실현하여 구조적 비용 문제를 근본적으로 해결한다."
    )),
    (25, 'heading', "개발 아이템 소개"),
    (27, 'body', (
        "마음온은 크게 2개의 모듈로 구성된 통합 정신건강 관리 생태계이다.\n"
        "① 사용자용 iOS 앱 — AI 프라이빗 다이어리: On-Device sLLM(Gemma-2-2b, 4-bit 양자화)을 활용한 "
        "24시간 오프라인 AI 감정 상담. 'AI 마음 친구 — 마음이(Maum-i)' 페르소나를 통한 공감적 대화와 "
        "다이어리 인터뷰 시퀀스(Event → Emotion → Reflection → Validation) 제공. 마음 온도(Mind Temperature) "
        "메트릭과 5단계 감정 매트릭스 기반 직관적 감정 시각화.\n"
        "② 의료진용 웹 포털 — 임상 대시보드: Vue.js 기반 의료진 전용 환자 모니터링 대시보드. "
        "암호화된 감정 트렌드, 수면 패턴, 위기 감지(Red Flag Alert) 실시간 수신. "
        "환자별 장·단기 AI 분석 리포트 및 기관 내 접근 통제 시스템."
    )),

    # ── 2. 실현 가능성 (Solution) ──
    (53, 'heading', "현재 개발 완료 현황 (MVP 단계 달성)"),
    (55, 'body', (
        "마음온은 이미 프로토타입을 넘어 MVP(Minimum Viable Product) 단계를 달성하였다.\n"
        "- iOS 앱(SwiftUI): Swift, MLX, CoreML 기반 MVP 완료 (App Bundle 모델 내장 포함)\n"
        "- 의료진 웹 포털: Vue.js 3, Vite, Pinia, Chart.js 기반 MVP 완료\n"
        "- 메인 백엔드(Hub): Django, DRF, MariaDB 기반 운영 중\n"
        "- AI 백엔드(Satellite): Flask, MongoDB, Gunicorn 기반 운영 중\n"
        "- On-Device AI 모델: Gemma-2-2b, 4-bit 양자화(MLX) 파인튜닝 완료 (~1.4GB)\n"
        "- 서버 AI 모델: Llama-3-8B (vLLM + RunPod Serverless) 연동 완료\n"
        "- 데이터 암호화: Fernet 대칭키 암호화 적용 완료\n"
        "- Hub↔Satellite 동기화: Push-on-Save Relay 구현 완료"
    )),
    (57, 'body', (
        "협약기간 내 구체화 계획:\n"
        "1. App Store 정식 출시 준비: UI/UX 고도화, QA 테스트, Apple 심사 대응\n"
        "2. 임상 파일럿 프로그램 운영: 2개 이상 기관(대학병원/정신건강복지센터)과 파일럿 MOU 체결\n"
        "3. On-Device AI 모델 v3 고도화: 한국어 감정 특화 코퍼스 확대 학습, 위기 감지 정확도 향상\n"
        "4. 보안 인증 취득: 개인정보보호 관련 인증(PIMS 등) 절차 착수"
    )),

    # ── 차별성 및 경쟁력 / 비즈니스 모델 ──
    (77, 'heading', "핵심 차별화 요소 (경쟁사 대비)"),
    (79, 'body', (
        "- 시장: 마인디(B2C 웰니스) vs 마음온(B2B/B2G 임상) → 기관 특화, 안정적 매출\n"
        "- 데이터 보안: 마인디(중앙 클라우드) vs 마음온(On-Device AI + 병원 로컬) → 물리적 데이터 주권\n"
        "- 의료 연동: 마인디(없음) vs 마음온(의료진 대시보드 + Red Flag 알림) → 임상적 활용도\n"
        "- 수익 모델: 마인디(개인 구독) vs 마음온(기관 SaaS 라이선싱) → 안정적 수익"
    )),
    (81, 'body', (
        "기술적 경쟁력:\n"
        "1. Network-OFF, AI-ON: 인터넷 연결 없이 24시간 AI 상담 가능\n"
        "2. Zero Token Cost: On-Device 실행으로 API 토큰 비용 완전 제거\n"
        "3. Hybrid AI Safety 아키텍처: 2단계 안전 감지 체계(정규식 + LLM 심층 분석)\n"
        "4. Fast-Load 패턴: AI 모델을 App Bundle에 사전 내장하여 앱 설치 즉시 AI 기능 활성화"
    )),
    (83, 'heading', "비즈니스 모델 (수익화 모델)"),
    (85, 'body', (
        "- B2G 기관 라이선스: 정신건강복지센터/보건소 대상, 환자 1인당 월 5,000~10,000원\n"
        "- B2B 기업 솔루션: 기업 EAP(직원지원프로그램) 대상, 직원 1인당 연 50,000원\n"
        "- B2B SDK 라이선싱: 3rd Party 앱 통합, On-Device AI SDK 라이선스\n"
        "- 프리미엄 사용자(B2C 보조): 개인 월 정기 구독, 월 9,900원\n"
        "핵심 강점: On-Device 구조로 사용자 수 증가가 서버 비용 증가로 이어지지 않으며, "
        "기관 라이선스의 높은 갱신율(Lock-in 효과)로 안정적 매출 구조 확보"
    )),
    (87, 'body', (
        "투자유치(자금확보) 전략:\n"
        "- Seed(2025 하반기, 1~2억 원): 정부 지원사업 + 엔젤 투자 → MVP 고도화, 임상 파일럿\n"
        "- Pre-A(2026 상반기, 5~10억 원): 디지털 헬스케어 전문 VC → DTx 인증 추진, 팀 확대\n"
        "- Series A(2027, 30억 이상): 임팩트 투자 + 전략적 투자 → 전국 B2G 확대, 해외 진출"
    )),

    # ── 4. 팀 구성 (Team) ──
    (106, 'heading', "대표자 보유 역량"),
    (108, 'body', (
        "기술 역량 — 풀스택 AI 엔지니어링:\n"
        "- 프론트엔드: Vue.js 3(Vite, Pinia), SwiftUI, Chart.js — 웹 대시보드 및 iOS 앱 독자 개발 완수\n"
        "- 백엔드: Python(Django/DRF, Flask), MariaDB, MongoDB — Hub/Satellite 이중 노드 아키텍처\n"
        "- AI: LLM 양자화(4-bit), 파인튜닝(LoRA), Apple MLX 최적화, vLLM 서빙, RunPod Serverless\n"
        "- 인프라: OCI 서버 운영, Nginx, Systemd, Git LFS, CI/CD"
    )),
    (110, 'body', (
        "핵심 성과:\n"
        "- 마음온 플랫폼 전체(iOS 앱 + 웹 대시보드 + 백엔드 + AI 모델) 1인 독자 개발 완수\n"
        "- On-Device AI 모델 파인튜닝 및 App Bundle 내장(Fast-Load 패턴) 구현\n"
        "- Hub/Satellite Push-on-Save 실시간 동기화 아키텍처 설계 및 운영"
    )),
]

for pidx, ptype, content in para_contents:
    if pidx >= len(doc.paragraphs):
        continue
    para = doc.paragraphs[pidx]
    if ptype == 'heading':
        set_para_heading(para, content)
    else:
        set_para_body(para, content)

print("✅ 본문 단락 완료")


# ═══════════════════════════════════════════════════════
# TABLE 4: 사업추진 일정(협약기간 내)
# ═══════════════════════════════════════════════════════
t4 = doc.tables[4]
t4_data = [
    ("1", "UX/UI 디자이너 채용", "25.04 ~ 25.05", "모바일 앱 디자인 전문 인력 1명"),
    ("2", "프론트엔드 개발자 채용", "25.04 ~ 25.06", "Vue.js/SwiftUI 개발 인력 1명"),
    ("3", "iOS 앱 UI/UX 고도화", "25.05 ~ 25.07", "디자인 시스템 정립, 사용성 테스트"),
    ("4", "On-Device AI 모델 v3 학습", "25.05 ~ 25.08", "한국어 감정 코퍼스 확대, 정확도 개선"),
    ("5", "시제품(최종 버전) 완성", "협약기간 말", "파일럿 피드백 반영 최종 버전"),
]
for ri, (n, c, p, d) in enumerate(t4_data):
    row = t4.rows[ri + 1]
    set_cell_text(row.cells[0], n)
    set_cell_text(row.cells[1], c)
    set_cell_text(row.cells[2], p)
    set_cell_text(row.cells[3], d)

print("✅ TABLE 4 완료")


# ═══════════════════════════════════════════════════════
# TABLE 5: 1단계 정부지원사업비 (20백만원)
# ═══════════════════════════════════════════════════════
t5 = doc.tables[5]
t5_data = [
    ("인건비", "▪ UX/UI 디자이너 급여 (3개월×3,000,000원)", "9,000,000"),
    ("재료비", "▪ AI 학습용 GPU 클라우드 사용료 (RunPod, 3개월)", "3,000,000"),
    ("지급수수료", "▪ Apple Developer Program 연간 회비", "200,000"),
    ("지급수수료", "▪ 서버 인프라 비용 (OCI 클라우드, 6개월)", "1,800,000"),
    ("외주용역비", "▪ 한국어 감정 데이터셋 구축 용역 (라벨링)", "4,000,000"),
]
for ri, (cat, desc, amt) in enumerate(t5_data):
    row = t5.rows[ri + 1]
    set_cell_text(row.cells[0], cat)
    set_cell_text(row.cells[1], desc, font_size=9)
    set_cell_text(row.cells[2], amt)

# R5 (원래 "…" 행) → 테스트 디바이스
if len(t5.rows) > 5:
    set_cell_text(t5.rows[5].cells[0], "재료비")
    set_cell_text(t5.rows[5].cells[1], "▪ 테스트 디바이스 구입 (iPhone, 2대)", font_size=9)
    set_cell_text(t5.rows[5].cells[2], "2,000,000")

# 합계 행
set_cell_text(t5.rows[-1].cells[2], "20,000,000", bold=True)

print("✅ TABLE 5 완료")


# ═══════════════════════════════════════════════════════
# TABLE 6: 2단계 정부지원사업비 (40백만원)
# ═══════════════════════════════════════════════════════
t6 = doc.tables[6]
t6_data = [
    ("인건비", "▪ 프론트엔드 개발자 급여 (5개월×4,000,000원)", "20,000,000"),
    ("외주용역비", "▪ 보안 인증(PIMS) 컨설팅 및 심사 비용", "5,000,000"),
    ("외주용역비", "▪ 임상 파일럿 운영 지원 (데이터 분석, IRB)", "3,000,000"),
    ("재료비", "▪ AI 모델 v3 고급 학습 GPU 클라우드 (A100, 2개월)", "5,000,000"),
]
for ri, (cat, desc, amt) in enumerate(t6_data):
    row = t6.rows[ri + 1]
    set_cell_text(row.cells[0], cat)
    set_cell_text(row.cells[1], desc, font_size=9)
    set_cell_text(row.cells[2], amt)

# R5 (원래 "…" 행) → 마케팅 + 전시회 + 서버
if len(t6.rows) > 5:
    set_cell_text(t6.rows[5].cells[0], "지급수수료")
    set_cell_text(t6.rows[5].cells[1], "▪ 마케팅/홍보, 전시회 참가비, 서버 확장 비용", font_size=9)
    set_cell_text(t6.rows[5].cells[2], "7,000,000")

# 합계
set_cell_text(t6.rows[-1].cells[2], "40,000,000", bold=True)

print("✅ TABLE 6 완료")


# ═══════════════════════════════════════════════════════
# TABLE 7: 사업추진 일정(전체 사업단계)
# ═══════════════════════════════════════════════════════
t7 = doc.tables[7]
t7_data = [
    ("1", "MVP 고도화 및 팀 구성", "2025 상반기", "UI/UX 개선, 핵심 인력 채용, AI 모델 v3 학습"),
    ("2", "App Store 정식 출시", "2025 3분기", "Apple 심사 대응 및 정식 출시"),
    ("3", "임상 파일럿 운영", "25.07 ~ 25.10", "대학병원/정신건강복지센터 2곳 이상 파일럿"),
    ("4", "B2G 사업 확장", "2026 상반기", "전국 정신건강복지센터 10곳 이상 도입"),
    ("5", "DTx 인증 및 해외 진출", "2027~", "디지털 치료제 인증, 동남아 시장 진출"),
]
for ri, (n, c, p, d) in enumerate(t7_data):
    row = t7.rows[ri + 1]
    if ri + 1 >= len(t7.rows):
        break
    set_cell_text(row.cells[0], n)
    set_cell_text(row.cells[1], c)
    set_cell_text(row.cells[2], p)
    set_cell_text(row.cells[3], d)

# "…" 행 비우기
if len(t7.rows) > 5:
    for ci in range(4):
        set_cell_text(t7.rows[5].cells[ci], "")

print("✅ TABLE 7 완료")


# ═══════════════════════════════════════════════════════
# TABLE 8: 팀 구성(안) — 4 rows (R0 헤더 + R1~R3)
# ═══════════════════════════════════════════════════════
t8 = doc.tables[8]
t8_data = [
    ("1", "대표 / CTO", "풀스택 개발 및 AI 엔지니어링 총괄",
     "컴퓨터공학 전공, 풀스택 개발(Vue.js, Swift, Python) 및 AI 엔지니어링(LLM 양자화, 파인튜닝) 역량 보유",
     "완료"),
    ("2", "자문위원", "임상 프로토콜 검증 및 자문",
     "임상심리학 박사, 정신건강의학 전문의",
     "완료"),
    ("3", "디자이너", "UX/UI 디자인",
     "시각디자인 학사, 모바일 앱 디자인 경력(1년 이상)",
     "예정('25.04)"),
]
for ri, (n, pos, task, skill, status) in enumerate(t8_data):
    row = t8.rows[ri + 1]
    set_cell_text(row.cells[0], n)
    set_cell_text(row.cells[1], pos)
    set_cell_text(row.cells[2], task, font_size=9)
    set_cell_text(row.cells[3], skill, font_size=9)
    set_cell_text(row.cells[4], status)

print("✅ TABLE 8 완료")


# ═══════════════════════════════════════════════════════
# TABLE 9: 협력 기관 — 4 rows (R0 헤더 + R1~R3)
# ═══════════════════════════════════════════════════════
t9 = doc.tables[9]
t9_data = [
    ("1", "○○대학교병원 정신건강의학과", "임상 연구 인프라, 환자 코호트 보유",
     "임상 파일럿 공동 운영, 효과성 데이터 수집", "25.06"),
    ("2", "○○구 정신건강복지센터", "B2G 현장 운영 경험, 지역사회 네트워크",
     "B2G 파일럿 도입, 현장 피드백 수집", "25.07"),
    ("3", "RunPod", "Serverless GPU 인프라 제공",
     "AI 모델 학습 및 서버사이드 추론 인프라 지원", "25.04"),
]
for ri, (n, name, cap, plan, timing) in enumerate(t9_data):
    row = t9.rows[ri + 1]
    set_cell_text(row.cells[0], n)
    set_cell_text(row.cells[1], name, font_size=9)
    set_cell_text(row.cells[2], cap, font_size=9)
    set_cell_text(row.cells[3], plan, font_size=9)
    set_cell_text(row.cells[4], timing)

print("✅ TABLE 9 완료")


# ═══════════════════════════════════════════════════════
# 최종 저장
# ═══════════════════════════════════════════════════════
doc.save(DST)
print(f"\n🎉 사업계획서 생성 완료: {DST}")
print(f"   파일 위치: {DST}")
