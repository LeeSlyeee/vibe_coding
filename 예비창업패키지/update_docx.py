import docx
import sys

def update_budget_tables(doc_path, output_path):
    print(f"Reading {doc_path}...")
    doc = docx.Document(doc_path)
    
    modified = False
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) >= 3:
            header_text = table.rows[0].cells[0].text.replace(' ', '')
            if '비목' in header_text or '비\t목' in header_text or '비\n목' in header_text:
                for row in table.rows:
                    cell_1_text = row.cells[1].text
                    
                    if "AI/ML 엔지니어" in cell_1_text:
                        print("Found AI/ML 엔지니어, updating to 인건비")
                        row.cells[0].paragraphs[0].text = "인건비"
                        modified = True
                        
                    if "GPU 클라우드" in cell_1_text:
                        print("Found GPU 클라우드, updating to 지급수수료")
                        row.cells[0].paragraphs[0].text = "지급수수료"
                        modified = True
                        
    if modified:
        doc.save(output_path)
        print(f"File saved to {output_path}")
    else:
        print("No modifications made.")

if __name__ == "__main__":
    if len(sys.argv) == 3:
        update_budget_tables(sys.argv[1], sys.argv[2])
    else:
        print("Usage: python update_docx.py <input> <output>")
