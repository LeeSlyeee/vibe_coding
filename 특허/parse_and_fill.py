import docx
from bs4 import BeautifulSoup
import re

html_path = '특허명세서_v2_마음온_AI_하이브리드_심리케어시스템.html'
template_path = '특허MS워드템플릿.docx'
out_path = '특허출원명세서_마음온_AI.docx'

with open(html_path, 'r', encoding='utf-8') as f:
    soup = BeautifulSoup(f, 'html.parser')

# Extract sections
# We'll build a dictionary mapping tags to text blocks
content_dict = {}

# A function to extract text after a heading
def get_text_for_heading(heading_text):
    heading = soup.find(lambda tag: tag.name in ['h1', 'h2', 'h3'] and heading_text in tag.get_text())
    if not heading:
        return []
    
    texts = []
    curr = heading.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2', 'h3']:
        # If it's a paragraph or list, extract
        if curr.name in ['p', 'ul', 'ol', 'div']:
            text = curr.get_text(separator='\n', strip=True)
            if text:
                texts.append(text)
        curr = curr.find_next_sibling()
    return texts

content_dict['【발명의 명칭】'] = get_text_for_heading('【발명의 명칭】')
content_dict['【기술분야】'] = get_text_for_heading('【기술분야】')
content_dict['【발명의 배경이 되는 기술】'] = get_text_for_heading('【배경기술】')

# Extract 선행기술
def get_prior_art():
    h3 = soup.find('h3', string=re.compile(r'선행기술문헌'))
    if not h3: return []
    texts = []
    curr = h3.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2', 'h3']:
        text = curr.get_text(separator='\n', strip=True)
        if text and '특허문헌' not in text:
            texts.append(text)
        curr = curr.find_next_sibling()
    return texts

content_dict['(【특허문헌】)'] = get_prior_art()

content_dict['【해결하고자 하는 과제】'] = get_text_for_heading('해결하고자 하는 과제')
content_dict['【과제의 해결 수단】'] = get_text_for_heading('과제의 해결 수단')
content_dict['【발명의 효과】'] = get_text_for_heading('발명의 효과')
content_dict['【도면의 간단한 설명】'] = get_text_for_heading('도면의 간단한 설명')
content_dict['【발명을 실시하기 위한 구체적인 내용】'] = get_text_for_heading('발명을 실시하기 위한 구체적인 내용')
content_dict['【요약】'] = get_text_for_heading('【요 약】')

# Claims
claims_h2 = soup.find('h2', string=re.compile(r'특허청구의 범위|청구범위'))
if claims_h2:
    curr = claims_h2.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2']:
        if curr.name == 'div' and 'claim' in curr.get('class', []):
            claim_text = curr.get_text(separator='\n', strip=True)
            # e.g., 【청구항 1】 \n text...
            match = re.match(r'(【청구항 \d+】)\s*(.*)', claim_text, re.DOTALL)
            if match:
                content_dict[match.group(1)] = [match.group(2)]
        curr = curr.find_next_sibling()

# Open DOCX
doc = docx.Document(template_path)

# Insert text after matching paragraphs
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text in content_dict:
        lines = content_dict[text]
        for line in reversed(lines):  # Reversed because we are injecting *after* the current paragraph
            # We want to add paragraphs after the matched heading
            p = paragraph.insert_paragraph_before(line)
            # Wait, insert_paragraph_before on the *next* sibling?
            # docx has no insert_after natively. 
            pass

# Better approach for python-docx: recreate the document or replace
new_doc = docx.Document()
# copy styles or just build it
for paragraph in doc.paragraphs:
    p = new_doc.add_paragraph()
    # copy alignment, style
    p.style = paragraph.style
    p.alignment = paragraph.alignment
    # copy runs
    for run in paragraph.runs:
        new_run = p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
    
    text = paragraph.text.strip()
    if text in content_dict:
        lines = content_dict[text]
        for line in lines:
            new_doc.add_paragraph(line)

new_doc.save(out_path)
print("Done")
