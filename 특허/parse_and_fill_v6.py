import docx
from bs4 import BeautifulSoup
import re
import os
from docx.shared import Pt, Inches

html_path = '특허명세서_v2_마음온_AI_하이브리드_심리케어시스템.html'
template_path = '특허MS워드템플릿.docx'
out_path = '특허출원명세서_마음온_AI.docx'

with open(html_path, 'r', encoding='utf-8') as f:
    soup = BeautifulSoup(f, 'html.parser')

def get_text_for_heading(heading_text):
    heading = soup.find(lambda tag: tag.name in ['h1', 'h2', 'h3'] and heading_text in tag.get_text())
    if not heading:
        return []
    texts = []
    curr = heading.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2', 'h3']:
        if curr.name in ['p', 'ul', 'ol', 'div']:
            # Replace <br> with newline to preserve line breaks
            for br in curr.find_all('br'):
                br.replace_with('\n')
            
            raw_text = curr.get_text()
            for line in raw_text.split('\n'):
                line = line.strip()
                if line:
                    texts.append(line)
        curr = curr.find_next_sibling()
    return texts

content_dict = {}

content_dict['【발명의 명칭】'] = get_text_for_heading('【발명의 명칭】')
content_dict['【기술분야】'] = get_text_for_heading('【기술분야】')
content_dict['【발명의 배경이 되는 기술】'] = get_text_for_heading('【배경기술】')

def get_prior_art():
    h3 = soup.find('h3', string=re.compile(r'선행기술문헌'))
    if not h3: return []
    texts = []
    curr = h3.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2', 'h3']:
        for br in curr.find_all('br'):
            br.replace_with('\n')
        raw_text = curr.get_text()
        for line in raw_text.split('\n'):
            line = line.strip()
            if line and '특허문헌' not in line or line.startswith('('):
                line = line.replace('【특허문헌】', '').replace('(특허문헌', '(특허문헌')
                if line.startswith('('):
                    texts.append(line)
        curr = curr.find_next_sibling()
    return texts

content_dict['(【특허문헌】)'] = get_prior_art()
content_dict['【해결하고자 하는 과제】'] = get_text_for_heading('해결하고자 하는 과제')
content_dict['【과제의 해결 수단】'] = get_text_for_heading('과제의 해결 수단')
content_dict['【발명의 효과】'] = get_text_for_heading('발명의 효과')
content_dict['【도면의 간단한 설명】'] = get_text_for_heading('도면의 간단한 설명')
content_dict['【발명을 실시하기 위한 구체적인 내용】'] = get_text_for_heading('발명을 실시하기 위한 구체적인 내용')
content_dict['【요약】'] = get_text_for_heading('【요약】')

# Process Claims
claims = []
claims_h2 = soup.find('h2', string=re.compile(r'특허청구의 범위|청구범위|특허청구범위'))
if claims_h2:
    curr = claims_h2.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2']:
        if curr.name == 'div' and 'claim' in curr.get('class', []):
            claim_text_div = curr.find_next_sibling('div', class_='claim-text')
            if claim_text_div:
                claim_num = curr.get_text(separator=' ', strip=True)
                
                for br in claim_text_div.find_all('br'):
                    br.replace_with('\n')
                
                raw_body = claim_text_div.get_text()
                body_lines = [line.strip() for line in raw_body.split('\n') if line.strip()]
                claims.append((claim_num, body_lines))
        curr = curr.find_next_sibling()

# Image mapping
image_base_path = '도면'
images_map = {
    '【도 1】': os.path.join(image_base_path, '도1_전체프로세스흐름도.png'),
    '【도 2】': os.path.join(image_base_path, '도2_동적질문트리알고리즘.png'),
    '【도 3】': os.path.join(image_base_path, '도3_하이브리드라우팅로직.png'),
    '【도 4】': os.path.join(image_base_path, '도4_메모리파기시퀀스.png'),
    '【도 5】': os.path.join(image_base_path, '도5_시스템블록도.png')
}

doc = docx.Document(template_path)
new_doc = docx.Document()

for paragraph in doc.paragraphs:
    p = new_doc.add_paragraph()
    p.style = paragraph.style
    p.alignment = paragraph.alignment
    for run in paragraph.runs:
        new_run = p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
    
    text_nospace = paragraph.text.replace(' ', '').strip()
    raw_text = paragraph.text.strip()

    if text_nospace == "【청구항1】":
        p.clear()
        
        for claim_num, body_lines in claims:
            # Clean up (독립 장치항) etc from the title for formal submission
            c_num_clean = claim_num.split('】')[0] + '】'
            cp = new_doc.add_paragraph()
            cp.add_run(c_num_clean).bold = True
            
            for index, line in enumerate(body_lines):
                line_p = new_doc.add_paragraph(line)
                # First line is usually preamble "~~ 시스템에 있어서,"
                # Following lines are elements ending with ";" so let's indent them
                if index > 0:
                    line_format = line_p.paragraph_format
                    line_format.left_indent = Inches(0.2)
        continue
    
    if re.match(r'【청구항[ \d]+】', text_nospace) and text_nospace != "【청구항1】":
        p.clear()
        continue
    if re.match(r'\(.*청구항.*\)', text_nospace):
        p.clear()
        continue

    # Map text content
    if raw_text in content_dict:
        lines = content_dict[raw_text]
        for line in lines:
            np = new_doc.add_paragraph(line)

    # Map images
    matched_img_key = None
    for img_key in images_map:
        if text_nospace == img_key.replace(' ', ''):
            matched_img_key = img_key
            break
            
    if matched_img_key:
        img_path = images_map[matched_img_key]
        if os.path.exists(img_path):
            img_p = new_doc.add_paragraph()
            img_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_p.add_run()
            img_run.add_picture(img_path, width=docx.shared.Inches(5.0))

new_doc.save(out_path)
print("Docx generation with properly indented multiline claims completed.")
