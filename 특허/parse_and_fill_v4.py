import docx
from bs4 import BeautifulSoup
import re

html_path = '특허명세서_v2_마음온_AI_하이브리드_심리케어시스템.html'
template_path = '특허MS워드템플릿.docx'
out_path = '특허출원명세서_마음온_AI.docx'

with open(html_path, 'r', encoding='utf-8') as f:
    soup = BeautifulSoup(f, 'html.parser')

def get_text_for_heading(heading_text):
    heading = soup.find(lambda tag: tag.name in ['h1', 'h2', 'h3'] and heading_text in tag.get_text())
    if not heading:
        return []
    texts = []
    curr = heading.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2', 'h3']:
        if curr.name in ['p', 'ul', 'ol', 'div']:
            text = curr.get_text(separator=' ', strip=True) # use ' ' separator to keep paragraphs intact
            if text:
                texts.append(text)
        curr = curr.find_next_sibling()
    return texts

content_dict = {}

content_dict['【발명의 명칭】'] = get_text_for_heading('【발명의 명칭】')
content_dict['【기술분야】'] = get_text_for_heading('【기술분야】')
content_dict['【발명의 배경이 되는 기술】'] = get_text_for_heading('【배경기술】')

def get_prior_art():
    h3 = soup.find('h3', string=re.compile(r'선행기술문헌'))
    if not h3: return []
    texts = []
    curr = h3.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2', 'h3']:
        # Extract <br> separated text lines individually
        for content in curr.stripped_strings:
            text = content.strip()
            if text and '특허문헌' not in text or text.startswith('('):
                text = text.replace('【특허문헌】', '').replace('(특허문헌', '(특허문헌') # Clean up
                if text.startswith('('):
                    texts.append(text)
        curr = curr.find_next_sibling()
    return texts

content_dict['(【특허문헌】)'] = get_prior_art()
content_dict['【해결하고자 하는 과제】'] = get_text_for_heading('해결하고자 하는 과제')
content_dict['【과제의 해결 수단】'] = get_text_for_heading('과제의 해결 수단')
content_dict['【발명의 효과】'] = get_text_for_heading('발명의 효과')
content_dict['【도면의 간단한 설명】'] = get_text_for_heading('도면의 간단한 설명')
content_dict['【발명을 실시하기 위한 구체적인 내용】'] = get_text_for_heading('발명을 실시하기 위한 구체적인 내용')
content_dict['【요약】'] = get_text_for_heading('【요약】') # Fixed space!

# Process Claims
claims = []
claims_h2 = soup.find('h2', string=re.compile(r'특허청구의 범위|청구범위|특허청구범위'))
if claims_h2:
    curr = claims_h2.find_next_sibling()
    while curr and curr.name not in ['h1', 'h2']:
        if curr.name == 'div' and 'claim' in curr.get('class', []):
            # The next element should be the claim-text
            claim_text_div = curr.find_next_sibling('div', class_='claim-text')
            if claim_text_div:
                claim_num = curr.get_text(separator=' ', strip=True) # e.g. 【청구항 1】
                claim_body = claim_text_div.get_text(separator=' ', strip=True)
                claims.append(f"{claim_num}\n{claim_body}")
        curr = curr.find_next_sibling()

doc = docx.Document(template_path)
new_doc = docx.Document()

# Read the template and process
for paragraph in doc.paragraphs:
    # Retain the exact original run layout so styles are preserved
    p = new_doc.add_paragraph()
    p.style = paragraph.style
    p.alignment = paragraph.alignment
    for run in paragraph.runs:
        new_run = p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
    
    text_nospace = paragraph.text.replace(' ', '').strip()
    raw_text = paragraph.text.strip()

    if text_nospace == "【청구항1】":
        p.clear()
        for claim in claims:
            match = re.match(r'(【청구항 \d+】.*?)\n(.*)', claim, re.DOTALL)
            if match:
                cp = new_doc.add_paragraph()
                cp.add_run(match.group(1)).bold = True
                new_doc.add_paragraph(match.group(2))
            else:
                new_doc.add_paragraph(claim)
        continue
    
    # We do not want to duplicate tags if they're in the template but we provided them already
    if re.match(r'【청구항[ \d]+】', text_nospace) and text_nospace != "【청구항1】":
        p.clear()
        continue
    if re.match(r'\(.*청구항.*\)', text_nospace): # like (독립 방법항)
        p.clear()
        continue

    # Try mapping using raw text
    for key in content_dict:
        if raw_text == key:
            lines = content_dict[key]
            for line in lines:
                new_doc.add_paragraph(line)

new_doc.save(out_path)
print("Finished adding everything to", out_path)
